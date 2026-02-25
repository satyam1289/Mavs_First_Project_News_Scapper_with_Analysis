import os
import streamlit as st
import pandas as pd
from io import BytesIO
import asyncio
from datetime import datetime, timedelta
import time
import auth  # Import our new auth module

# app2.py
from advanced_ner_extractor import load_ner_model, analyze_specific_brands
# Import our helper tools (which we wrote in other files)
from gdelt_fetcher import fetch_gdelt_simple, WHITELIST_DOMAINS
from article_scraper import enhance_articles_async
from sector_classifier import classify_sector
import db_manager

# --- PREDEFINED OTHERS BRAND POOL ---
OTHER_BRANDS_POOL = [
    "TCS", "Infosys", "Wipro", "HCLTech", "IBM", "Deloitte", "PwC", "EY", "KPMG", 
    "LTIMindtree", "Persistent Systems", "Zensar Technologies", "Happiest Minds", "Mu Sigma", 
    "DataRobot", "C3.ai", "Palantir Technologies", "Scale AI", "LeewayHertz", "Elysium Technologies", 
    "Ksolves", "Innowise Group", "ScienceSoft", "Addepto", "Markovate", "Sarvika Technologies", 
    "Intellectyx", "Millipixels Interactive", "InnovationM", "Sigmoid", "Google DeepMind", 
    "Meta AI", "Mistral AI", "Cohere", "xAI", "Perplexity AI", "DeepSeek", "Stability AI", 
    "Inflection AI", "Hugging Face", "Character.ai", "Adept AI", "Reka AI", "01.AI", 
    "LightOn", "Aleph Alpha", "AI21 Labs", "Kyutai", "Suno", "ElevenLabs", "AMD", "Intel", 
    "TSMC", "Broadcom", "Qualcomm", "Arm Holdings", "Groq", "Cerebras Systems", "SambaNova Systems", 
    "Graphcore", "Tenstorrent", "Lightmatter", "CoreWeave", "Lambda Labs", "ASML", 
    "Marvell Technology", "Micron Technology", "SK Hynix", "Supermicro", "SiliconFlow", 
    "Databricks", "Snowflake", "Runway", "Pika Labs", "Luma AI", "Sora", "Cursor", 
    "Cognition AI", "Harvey", "Abnormal Security", "Glean", "Shield AI", "Anduril Industries", 
    "Waymo", "Tesla", "Canva", "Adobe", "Notion", "Jasper", "Copy.ai", "Synthesia", 
    "Midjourney", "Haptik", "Uniphore", "Arya.ai", "Mad Street Den", "Locus.sh", "SigTuple", 
    "Ambience Healthcare", "Safe Superintelligence", "Sarvam AI", "Krutrim", "Tata Elxsi", 
    "Netweb Technologies", "E2E Networks", "Reliance Jio (JioBrain)", "Qure.ai", "Yellow.ai", 
    "GreyOrange", "Cropin", "Bhashini", "BharatGen", "CoRover", "Gnani.ai", "Entropik", 
    "Skit.ai", "Niramai", "Intello Labs", "Myelin Foundry", "Rephrase.ai", "Observe.AI", 
    "LogiNext", "Assert AI", "Kore.ai", "Active.ai", "Vernacular.ai", "Staqu", 
    "AIndra Systems", "Soket AI Labs"
]

INDIAN_SOURCES_KEYWORDS = [
    "economic times", "times of india", "the times of india", "toi", "timesofindia", 
    "the hindu", "hindustantimes", "hindustan times", "the hindustan times",
    "business standard", "livemint", "mint", "ndtv", "india today", "firstpost",
    "indian express", "financial express", "moneycontrol", "yourstory", "inc42",
    "business today", "zee news", "news18", "deccan herald", "outlook india",
    "the week", "businessworld", "forbes india", "fortune india", "vccircle",
    "techcircle", "entrackr", "trak.in", "techgig", "analytics india", "cnbc tv18",
    "pti", "ians", "ani", "deccan chronicle", "pib", "dainik", "jagran", "amar ujala",
    "bhaskar", "lokmat", "malayala", "mathrubhumi", "enadu", "sakshi", "dinakaran",
    "the pioneer", "tribune india", "new indian express", "telangana today",
    "kashmir reader", "greater kashmir", "assam tribune", "sentinel assam"
]

def is_indian_source(source_name: str) -> bool:
    if not source_name: return False
    sn = source_name.lower()
    # Check if any keyword matches
    for kw in INDIAN_SOURCES_KEYWORDS:
        if kw in sn:
            return True
    return False

# --- PAGE SETUP ---
st.set_page_config(page_title="News Intelligence", layout="wide", initial_sidebar_state="collapsed")

# --- CUSTOM STYLING (CSS) ---
st.markdown("""
<style>
    /* #MainMenu {visibility: hidden;} */
    /* header {visibility: hidden;} */
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
</style>
""", unsafe_allow_html=True)

# --- THEME CONTROL ---
if 'theme' not in st.session_state:
    st.session_state.theme = 'dark'

def apply_theme():
    if st.session_state.theme == 'light':
        st.markdown("""
        <style>
            /* 1. Global App & Sidebar Backgrounds */
            .stApp { background-color: #FFFFFF !important; color: #000000 !important; }
            [data-testid="stSidebar"] { background-color: #f8f9fa !important; border-right: 1px solid #eeeeee !important; }
            [data-testid="stSidebar"] * { color: #000000 !important; }
            section[data-testid="stSidebar"] .stMarkdown { color: #000000 !important; }
            
            /* 2. Typography & Core Elements */
            .stMarkdown, .stText, h1, h2, h3, h4, p, span, li, label, div { color: inherit; }
            .stMarkdown, .stText, h1, h2, h3, h4, p, span, li, label, .stCaption { color: #000000 !important; }
            header { background-color: rgba(255,255,255,0.8) !important; backdrop-filter: blur(10px); }
            
            /* 3. Metric Styling - Enforce Uniform Height and Centering */
            [data-testid="stMetric"] { 
                background-color: #ffffff !important; 
                border: 1px solid #e0e0e0 !important; 
                border-radius: 12px !important; 
                padding: 15px !important; 
                box-shadow: 0 4px 12px rgba(0,0,0,0.05) !important;
                min-height: 120px !important;
                display: flex !important;
                flex-direction: column !important;
                justify-content: center !important;
                align-items: center !important;
                text-align: center !important;
            }
            [data-testid="stMetricValue"] { color: #000000 !important; font-weight: 700 !important; font-size: 1.8rem !important; }
            [data-testid="stMetricLabel"] { color: #555555 !important; font-size: 1rem !important; margin-bottom: 5px !important; }
            
            /* 4. Dataframes & Tables */
            [data-testid="stDataFrame"], [data-testid="stTable"], .stDataFrame, .stTable { 
                background-color: #FFFFFF !important; 
                border: 1px solid #eeeeee !important;
                color: #000000 !important;
            }
            [data-testid="stDataFrame"] div, [data-testid="stDataFrame"] span, [data-testid="stDataFrame"] td { color: #000000 !important; }
            div[data-testid="stDataFrame"] > div { background-color: #FFFFFF !important; }
            
            /* 5. Inputs & Dropdowns (Broad Coverage) */
            [data-testid="stTextInput"] input, 
            [data-testid="stTextArea"] textarea,
            [data-testid="stNumberInput"] input,
            [data-testid="stDateInput"] input,
            [data-testid="stSelectbox"] div[data-baseweb="select"],
            [data-testid="stMultiSelect"] div[data-baseweb="select"],
            .stTextInput input, .stTextArea textarea {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                border: 1px solid #cccccc !important;
                border-radius: 8px !important;
                -webkit-text-fill-color: #000000 !important;
            }
            /* Autofill Fix */
            input:-webkit-autofill,
            input:-webkit-autofill:hover, 
            input:-webkit-autofill:focus, 
            input:-webkit-autofill:active {
                -webkit-box-shadow: 0 0 0 50px white inset !important;
                -webkit-text-fill-color: black !important;
            }
            /* Input Containers */
            [data-testid="stTextInput"] > div, [data-testid="stTextArea"] > div {
                background-color: #FFFFFF !important;
            }
            /* Password Eye Icon Fix */
            [data-testid="stTextInput"] button {
                background-color: transparent !important;
                color: #000000 !important;
                border: none !important;
            }
            
            /* 6. Buttons - Standard & Primary */
            .stButton button, .stDownloadButton button {
                background-color: #ffffff !important;
                color: #000000 !important;
                border: 1px solid #cccccc !important;
                border-radius: 8px !important;
                transition: all 0.2s ease !important;
                font-weight: 500 !important;
            }
            
            /* Force Primary style on ALL form buttons and specified primary buttons */
            [kind="primary"], [data-testid="stForm"] button, .stButton button[kind="primary"] {
                background-color: #2c3e50 !important;
                color: #FFFFFF !important;
                border: none !important;
                font-weight: 600 !important;
            }
            [kind="primary"]:hover { background-color: #34495e !important; box-shadow: 0 4px 12px rgba(0,0,0,0.15) !important; }
            
            /* 7. Tabs & Navigation */
            .stTabs [data-baseweb="tab-list"] { background-color: transparent !important; border-bottom: 2px solid #eeeeee !important; }
            .stTabs [data-baseweb="tab"] { color: #888888 !important; padding: 10px 20px !important; }
            .stTabs [aria-selected="true"] { color: #000000 !important; font-weight: bold !important; border-bottom-color: #2c3e50 !important; }

            /* 8. Status & Messages (Outputs) */
            [data-testid="stNotification"], [data-testid="stStatusWidget"], .stAlert { 
                border-radius: 10px !important; 
                border: 1px solid rgba(0,0,0,0.05) !important; 
                background-color: #f8f9fa !important;
            }
            [data-testid="stNotification"] *, .stAlert * { color: #000000 !important; }
            
            /* 9. Expanders & Containers */
            .streamlit-expanderHeader { color: #000000 !important; background-color: #f8f9fa !important; border-radius: 8px !important; border: 1px solid #eeeeee !important; }
            .streamlit-expanderContent { background-color: #ffffff !important; border: 1px solid #eeeeee !important; border-top: none !important; border-radius: 0 0 8px 8px !important; }
            [data-testid="stForm"], [data-testid="stVerticalBlock"] > div.element-container:has(div.stContainer) { 
                border: 1px solid #eeeeee !important; 
                border-radius: 12px !important; 
                padding: 20px !important; 
                background-color: #ffffff !important; 
            }
            /* Vertical Scrollable Container */
            [data-testid="stVerticalBlockBorderWrapper"] { border-color: #eeeeee !important; }

            /* 10. Toggles & Checkboxes */
            [data-testid="stCheckbox"] label, [data-testid="stToggle"] label { color: #000000 !important; }

            /* Divider lines */
            hr { border-color: #eeeeee !important; }
            
            /* Scrollbars for a premium feel */
            ::-webkit-scrollbar { width: 8px; height: 8px; }
            ::-webkit-scrollbar-track { background: #f1f1f1; }
            ::-webkit-scrollbar-thumb { background: #ccc; border-radius: 10px; }
            ::-webkit-scrollbar-thumb:hover { background: #999; }
        </style>
        """, unsafe_allow_html=True)
    else:
        # Dark Mode Consistent Styling
        st.markdown("""
        <style>
            .stApp { background-color: #0e1117 !important; }
            [data-testid="stSidebar"] { background-color: #1a1c24 !important; border-right: 1px solid #2d313e !important; }
            
            [data-testid="stMetric"] { 
                background-color: #1a1c24 !important; 
                border: 1px solid #2d313e !important; 
                border-radius: 12px !important; 
                padding: 15px !important; 
                box-shadow: 0 4px 12px rgba(0,0,0,0.3) !important;
                min-height: 120px !important;
                display: flex !important;
                flex-direction: column !important;
                justify-content: center !important;
                align-items: center !important;
                text-align: center !important;
            }
            [data-testid="stMetricValue"] { color: #FFFFFF !important; font-weight: 700 !important; font-size: 1.8rem !important; }
            [data-testid="stMetricLabel"] { color: #a1a1a1 !important; font-size: 1rem !important; margin-bottom: 5px !important; }
            
            .streamlit-expanderHeader { background-color: #1a1c24 !important; border-radius: 8px !important; border: 1px solid #2d313e !important; }
            .stButton button, .stDownloadButton button { border-radius: 8px !important; transition: all 0.2s ease !important; }
            
            /* Inputs in Dark Mode */
            [data-testid="stTextInput"] input, 
            [data-testid="stTextArea"] textarea, 
            [data-testid="stSelectbox"] div[data-baseweb="select"],
            [data-Baseweb="tab"] {
                border-radius: 8px !important;
            }
            
            /* Scrollbars for Dark Mode */
            ::-webkit-scrollbar { width: 8px; height: 8px; }
            ::-webkit-scrollbar-track { background: #0e1117; }
            ::-webkit-scrollbar-thumb { background: #3d4455; border-radius: 10px; }
            ::-webkit-scrollbar-thumb:hover { background: #4d5565; }
        </style>
        """, unsafe_allow_html=True)

def get_chart_theme():
    return "plotly_white" if st.session_state.theme == "light" else "plotly_dark"

col_theme = st.columns([0.9, 0.1])
with col_theme[1]:
    theme_icon = "🌞" if st.session_state.theme == "dark" else "🌙"
    if st.button(theme_icon, help="Switch to " + ("Bright" if st.session_state.theme == "dark" else "Dark") + " Mode"):
        st.session_state.theme = 'light' if st.session_state.theme == 'dark' else 'dark'
        st.rerun()

apply_theme()

@st.cache_resource(show_spinner=False)
def get_ner_pipeline():
    """
    Cached function to load the NER model once across the app session.
    """
    model, available = load_ner_model()
    return model

@st.cache_resource(show_spinner=False)
def preload_sector_classifier():
    """
    Pre-load the SBERT model and embeddings in the background so 
    it's instantly available for the first search query.
    """
    import threading
    def _background_load():
        try:
            from sector_classifier import build_sector_embeddings
            print("🚀 Pre-loading SBERT embeddings in background thread...")
            build_sector_embeddings()
            print("✅ SBERT embeddings ready.")
        except Exception as e:
            print(f"⚠️ Failed to pre-load embeddings: {e}")
            pass
            
    t = threading.Thread(target=_background_load, daemon=True)
    t.start()
    return True

# Trigger pre-load silently
preload_sector_classifier()

# --- AUTHENTICATION FUNCTIONS ---

def login_page():
    st.title("🔐 Login")
    
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("Login", type="primary")
        
        if submit:
            if not username or not password:
                st.warning("⚠️ Please enter both username and password.")
            else:
                success, message = auth.authenticate_user(username, password)
                
                if success:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.success(f"✅ {message}!")
                    st.rerun()
                else:
                    if message == "User not found":
                        st.warning("⚠️ User not registered. Please switch to the **Sign Up** tab.")
                    elif message == "Incorrect password":
                        st.error("❌ Incorrect password. Please try again.")
                    else:
                        st.error(f"❌ {message}")

def signup_page():
    st.title("📝 Sign Up")
    
    with st.form("signup_form"):
        new_user = st.text_input("Username")
        new_pass = st.text_input("Password", type="password")
        confirm_pass = st.text_input("Confirm Password", type="password")
        submit = st.form_submit_button("Sign Up", type="primary")
        
        if submit:
            if new_pass != confirm_pass:
                st.error("Passwords do not match!")
            elif len(new_pass) < 6:
                st.error("Password must be at least 6 characters.")
            else:
                if auth.create_user(new_user, new_pass):
                    st.success(f"✅ Account created for **{new_user}**! Please login.")
                    st.info("👉 Go to the **Login** tab to sign in.")
                else:
                    st.error(f"⚠️ Username **'{new_user}'** already exists. Please choose a different one.")

def logout():
    st.session_state['logged_in'] = False
    st.session_state['username'] = None
    st.rerun()

# --- MAIN DASHBOARD (Wrapped) ---

import db_manager # Import database manager

# ... (imports) ...

# --- PAGE SETUP ---
st.set_page_config(page_title="News Intelligence", layout="wide", initial_sidebar_state="collapsed")

# ... (CSS and Theme logic remains same) ...

@st.cache_resource(show_spinner=False)
def get_ner_pipeline():
    # ... (remains same) ...
    model, available = load_ner_model()
    return model

# ... (Auth functions remain same) ...

# --- MAIN DASHBOARD (Wrapped) ---

def show_dashboard():
    # ... (Sidebar Logout remains same) ...
    with st.sidebar:
        st.write(f"Logged in as: **{st.session_state.get('username', 'User')}**")
        if st.button("Log Out"):
            logout()
        st.markdown("---")

    # --- APP HEADER ---
    if os.path.exists("Mavericks logo.png"):
        st.image("Mavericks logo.png", width=150)
    st.title("📰 News Search Engine")
    
    # DB STATS
    stats = db_manager.get_stats()
    st.caption(f"Database contains **{stats['total_articles']}** articles from **{stats['total_sources']}** sources.")

    # --- GPU STATUS INDICATOR ---
    with st.expander("🖥️ System Status (GPU)", expanded=False):
        import torch
        if torch.cuda.is_available():
            device_name = torch.cuda.get_device_name(0)
            mem_alloc = torch.cuda.memory_allocated(0) / 1024**2
            mem_res = torch.cuda.memory_reserved(0) / 1024**2
            
            st.success(f"✅ GPU Detected: **{device_name}**")
            st.write(f"**Memory Allocated:** {mem_alloc:.2f} MB")
            st.write(f"**Memory Reserved:** {mem_res:.2f} MB")
            
            if st.button("⚡ Force Load Models (Test GPU)"):
                with st.spinner("Loading AI Models into VRAM..."):
                    # Clear cache to actually force a reload
                    get_ner_pipeline.clear()
                    # Force init
                    get_ner_pipeline()
                    from sector_classifier import get_sbert_model
                    get_sbert_model()
                    st.rerun() # Refresh to show new memory stats
        else:
            st.error("❌ GPU Not Detected (Running on CPU)")

    # REMOVED: st.session_state.articles initialization
    # We now query the DB directly.

    tab_search, tab_analysis = st.tabs(["🔍 Find News Articles", "📊 Brand Analysis"])

    with tab_search:
        # --- INPUT SECTION (Search Bar) ---

        # Map friendly names to Google News codes
        REGION_MAP = {
        "India 🇮🇳": "IN:en",
        "USA 🇺🇸": "US:en",
        "UK 🇬🇧": "GB:en",
        "Australia 🇦🇺": "AU:en",
        "Canada 🇨🇦": "CA:en",
        "Singapore 🇸🇬": "SG:en",
        "New Zealand 🇳🇿": "NZ:en",
        "Ireland 🇮🇪": "IE:en",
        "South Africa 🇿🇦": "ZA:en",
        "Philippines 🇵🇭": "PH:en",
        "Malaysia 🇲🇾": "MY:en",
        "Pakistan 🇵🇰": "PK:en",
        "Hong Kong 🇭🇰": "HK:en",
        "UAE 🇦🇪": "AE:en",
        "Europe 🇪🇺": "EU:en",
        "Global 🌐": "WORLD:en"
    }

        # Layout: [Sector/Keyword (2)] [Region (1)] [Days (1)]
        col1, col2, col3 = st.columns([2, 1, 1])

        with col1:
            sector_input = st.selectbox(
                "📂 Select Sector", 
                ["Lifestyle", "Sustainability", "Technology", "Artificial Intelligence", "Health", "Finance", "Education", "Sports", "Startups", "CUSTOM"],
                index=2 # Default to Technology
            )
            
            if sector_input == "CUSTOM":
                custom_keyword = st.text_input("🔍 Enter Custom Sector/Keyword", help="Type your topic")
                query = custom_keyword
            else:
                query = sector_input

        with col2:
            # Moved from Sidebar to Main UI for visibility
            selected_region_names = st.multiselect(
                "🌍 Regions",
                options=list(REGION_MAP.keys()),
                default=["India 🇮🇳"],
                help="Select which countries to source news from."
            )
            
            # Convert names to codes
            selected_region_codes = [REGION_MAP[name] for name in selected_region_names]
            
            if not selected_region_codes:
                selected_region_codes = list(REGION_MAP.values())

        with col3:
            # OPTIMIZED: Date Range Picker for precision
            today = datetime.now().date()
            last_week = today - timedelta(days=7)
            date_range = st.date_input(
                "📅 Select Date Range",
                value=(last_week, today),
                min_value=today - timedelta(days=365*2),
                max_value=today,
                format="DD/MM/YYYY"
            )
            
            start_date = None
            end_date = None

            if len(date_range) == 2:
                start_date, end_date = date_range
                start_date = datetime.combine(start_date, datetime.min.time())
                end_date = datetime.combine(end_date, datetime.max.time())
            elif len(date_range) == 1:
                st.warning("Please select an End Date.")
                st.stop()
            else:
                st.warning("Please select a valid date range.")
                st.stop()

        st.markdown("---")

    # --- CONFIGURATION (Sidebar) ---
    with st.sidebar:
        st.header("⚙️ Advanced Settings")
        
        # We use keys here to ensure the state persists when other UI elements (like date picker) trigger a rerun.
        use_tor = st.toggle("🛡️ Use Tor Proxy", value=st.session_state.get('use_tor_toggle', False), key='use_tor_toggle', help="Route all search & scraping through Tor (127.0.0.1:9150) to avoid rate limits.")
        saturation_mode = st.toggle("🔥 Saturation Mode", value=st.session_state.get('saturation_mode_toggle', False), key='saturation_mode_toggle', help="Deep search mode to hit maximum article volume (slow!)")
        search_entire_web = st.toggle("🌐 Search Entire Web", value=st.session_state.get('search_entire_web_toggle', False), key='search_entire_web_toggle', help="Disable the domain whitelist and search the entire internet for this topic.")
        
        if use_tor:
            # Check if Tor is actually running
            import socket
            try:
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                sock.settimeout(1) # 1 second timeout
                result = sock.connect_ex(('127.0.0.1', 9150))
                sock.close()
                if result == 0:
                    st.info("🛡️ **Tor Proxy Enabled**: IP rotation will be used if rate limits are hit.")
                else:
                    st.error("⚠️ **Tor Error**: Port 9150 is closed. Is the Tor Browser running?")
            except Exception:
                 st.warning("⚠️ **Tor Check Failed**: Could not verify Tor connection.")
            
        max_articles_selection = st.selectbox(
            "🔢 Max Articles", 
            options=[10, 25, 50, 100, 200, 500, 1000, 2000, 5000, "All"],
            index=9, # Default to "All"
            help="Limit the number of articles to scrape to save time."
        )
        
        # Convert "All" to a large number (e.g. 50,000)
        if max_articles_selection == "All":
            max_articles_limit = 50000
        else:
            max_articles_limit = max_articles_selection
            
        st.markdown("---")
        st.subheader("📰 Source Publications")
        # Pre-fill with our curated list
        default_whitelist = ", ".join(WHITELIST_DOMAINS)
        whitelist_input = st.text_area(
            "Allowed Domains (Comma Separated)", 
            value=default_whitelist,
            height=150,
            help="Only articles from these domains will be searched. Add or remove domains as needed."
        )
        
        # Parse the input into a list
        custom_whitelist = [d.strip() for d in whitelist_input.split(",") if d.strip()]

        # --- SEARCH ACTION ---
        # This runs when you click the big red button
        st.markdown("<br>", unsafe_allow_html=True) # Spacer
        if st.button("🚀 Find News Articles", type="primary", use_container_width=True):
            # Initialize Progress Bar immediately here to show 0%
            main_progress = st.progress(0, text="0% complete - Initializing Omega Strategy...")
            # --- CUSTOM LOADER ---
            # Increased size: Ratio 2:5, Width 250
            col_img, col_txt = st.columns([2, 5])
            with col_img:
                if os.path.exists("loader.jpg"):
                    st.image("loader.jpg", width=250)
            with col_txt:
                st.markdown(f"### {'🛡️' if use_tor else '🚀'} {'Tor Deep Search' if use_tor else 'Resilient Deep Search'} Active...")
                st.markdown(f"We are conducting a deep, stable search. {'Routing through Tor for IP rotation.' if use_tor else 'Using traffic-smoothing delays.'} Please stay with us... ⏳")
                # Add a Stop button here
                if st.button("🛑 Stop & Save", use_container_width=True):
                    pass # This click triggers a RerunException catching loop below

            # --- INTERNAL CLASSIFICATION (For Custom Keywords) ---
            if sector_input == "CUSTOM" and query:
                # Get API key from secrets
                try:
                    gemini_key = st.secrets.get("general", {}).get("GEMINI_API_KEY")
                except:
                    gemini_key = None
                
                # Classify using Hybrid approach (Gemini > SBERT > Keywords)
                classified_sector = classify_sector(query, api_key=gemini_key)
                
                # Store for display later
                st.session_state.classified_sector = classified_sector
                print(f"DEBUG: Hybrid Classification for '{query}': {classified_sector}")

            else:
                # Reset if not custom
                st.session_state.classified_sector = None

            
            # --- PROGRESSIVE LOADING STATUS ---
            with st.status("🤖 AI Agent is working...", expanded=True) as status:
                
                # STEP 1: FIND LINKS
                # Fake a small progress update to show activity
                main_progress.progress(10, text="10% complete - Initializing Traffic Smoothing...")
                status.write(f"🔍 Mode: {'Tor Proxy' if use_tor else 'Standard'} | Scanning {start_date.date()} to {end_date.date()} for stable retrieval...")
                
                # Callback to update the UI during the Search Phase (10% -> 50%)
                def search_progress_handler(completed, total):
                    if total > 0:
                        pct = 10 + int((completed / total) * 40) # Scale 0-100 to 10-50
                        main_progress.progress(pct, text=f"Searching... Found {completed} links from {start_date.date()}...")
                
                # 1. FETCH
                # Pass explicit START and END dates
                raw_articles = fetch_gdelt_simple(
                    keyword=query,
                    start_date=start_date,
                    end_date=end_date,
                    max_articles=max_articles_limit,
                    progress_callback=search_progress_handler,
                    target_regions=selected_region_codes,
                    sector_context=st.session_state.get('classified_sector') if sector_input == "CUSTOM" else sector_input,
                    use_tor=use_tor,
                    saturation_mode=saturation_mode,
                    whitelist_override=custom_whitelist,
                    search_entire_web=search_entire_web
                )
                
                # Jump to 50% after finding links
                main_progress.progress(50, text=f"50% complete - Found {len(raw_articles)} links... Processing...")
                
                if not raw_articles:
                    status.update(label="❌ No news found!", state="error", expanded=False)
                    st.error("No news found for this keyword. Please try another.")
                    st.session_state.articles = []
                else:
                    status.write(f"✅ Found {len(raw_articles)} links from around the web.")
                    
                    # STEP 2: READ CONTENT
                    status.write(f"📖 Visiting all {len(raw_articles)} websites to extract content...")
                    
                    # This little function updates the main progress bar
                    def update_progress(current, total):
                        # We map the scraping progress (0-100%) to the remaining main progress (50-100%)
                        scrape_percent = (current / total)
                        total_percent = int(50 + (scrape_percent * 50))
                        
                        main_progress.progress(total_percent, text=f"{total_percent}% complete - Reading article {current}/{total}")
                        
                        # Update text every few items inside the status box too
                        if current % 10 == 0 or current == total:
                                status.update(label=f"📖 Reading articles... ({int(scrape_percent*100)}%)")
                    
                    # RUN THE SCRAPER! (This visits all sites)
                    try:
                        # DEBUG: Check raw articles count
                        st.write(f"DEBUG: Starting scrape for {len(raw_articles)} articles...")
                        
                        # STEP 2: SCRAPE CONTENT
                        # We now get an error report too!
                        enhanced_articles, error_report = asyncio.run(enhance_articles_async(
                            raw_articles, 
                            limit=max_articles_limit, 
                            progress_callback=update_progress,
                            use_tor=use_tor
                        ))
                        
                        # DEBUG: Check result count
                        st.write(f"DEBUG: Scrape finished. Enhanced {len(enhanced_articles)} articles.")
                        
                        main_progress.progress(90, text="90% complete - Saving to Database...")
                        
                        # SAVE TO DATABASE
                        saved_count = db_manager.save_articles(enhanced_articles)
                        st.write(f"💾 Saved **{saved_count}** new articles to the database.")
                        
                        main_progress.progress(100, text="100% complete - Done!")
                        
                        st.session_state.last_query = query
                        
                        # Collapse the status box when done
                        status.update(label="✅ All Done! Articles ready.", state="complete", expanded=False)
                        
                        if error_report:
                            with st.expander("⚠️ Scraping Issues Report", expanded=False):
                                st.warning("Some articles could not be fully scraped. This is normal for protected sites.")
                                for domain, stats in error_report.items():
                                    reasons = ", ".join(list(stats['reasons'])[:3])
                                    st.write(f"**{domain}**: {stats['failed']}/{stats['total']} failed ({reasons})")
                                    
                    except Exception as e:
                        # Normal runtime errors
                        st.error(f"CRITICAL ERROR during scraping: {e}")
                        status.update(label="❌ Error during scraping", state="error")
                    except BaseException as e:
                        # This safely catches stream death (Streamlit StopException/RerunException)
                        if 'raw_articles' in locals() and raw_articles:
                            st.toast("🛑 Stopping and saving partial results...")
                            db_manager.save_articles(raw_articles)
                        raise e # re-raise to complete the Streamlit rerun

        # --- DISPLAY RESULTS (PAGINATED DB) ---
        st.markdown("---")
        st.subheader(f"📋 Results")
        
        # Filter Controls
        col_filter1, col_filter2 = st.columns([2, 2])
        with col_filter1:
            search_filter = st.text_input("🔍 Filter by Keyword", placeholder="Search in database...")
        with col_filter2:
            sector_filter = st.selectbox("📂 Filter by Sector", ["All", "Technology", "Artificial Intelligence", "Finance", "Health", "Sustainability", "Startups"])

    with tab_analysis:
        st.subheader("🔍 Competitive Analysis")
        st.caption("Compare the media exposure of specific brands or topics across the entire database.")
        
        ca_col1, ca_col2 = st.columns([3, 1])
        with ca_col1:
            target_brands_input = st.text_area("Target Brands/Topics (comma-separated)", placeholder="e.g. Reliance, Tata, Adani\nYou can enter as many brands as you need.", height=100)

        with ca_col2:
            st.markdown("<br>", unsafe_allow_html=True)
            analyze_btn = st.button("Analyze Exposure", use_container_width=True)
            
        if analyze_btn:
            if not target_brands_input:
                st.warning("Please enter a target brand to analyze.")
            else:
                target_list = [b.strip() for b in target_brands_input.split(",") if b.strip()]
                st.session_state.ca_brands_to_analyze = target_list
                
                with st.spinner("🚀 Intelligent Analysis in progress..."):
                    # 1. Prepare full list of brands to check (Targets + Pool)
                    full_list = list(target_list)
                    for b in OTHER_BRANDS_POOL:
                        if b not in full_list:
                            full_list.append(b)
                            
                    current_db_total = db_manager.get_total_count()
                    all_results = {}
                    missing_brands = []
                    
                    # 2. Check Cache
                    for brand in full_list:
                        cached_data, count_at_cache = db_manager.get_cached_analysis(brand)
                        if cached_data and count_at_cache == current_db_total:
                            all_results[brand] = cached_data
                        else:
                            missing_brands.append(brand)
                            
                    # 3. Run Analysis for Missing
                    if missing_brands:
                        analysis_articles = db_manager.get_articles(limit=1000000)
                        if analysis_articles:
                            # We analyze individual brands (no internal 'Others' grouping here)
                            new_results = analyze_specific_brands(analysis_articles, missing_brands)
                            for b, data in new_results.items():
                                all_results[b] = data
                                db_manager.save_analysis_cache(b, data, current_db_total)
                                
                    # 4. Final Aggregation
                    # final_analysis contains targets + one "Others" entry
                    final_analysis = {}
                    others_agg = {
                        "mentions": 0, "articles": 0, "sources": {}, "timeline": {}, 
                        "sentiment": {"Positive": 0, "Neutral": 0, "Negative": 0},
                        "article_samples": {"Positive": [], "Neutral": [], "Negative": []}
                    }
                    
                    target_set_lower = {b.lower().strip() for b in target_list}
                    
                    for brand, data in all_results.items():
                        b_clean = brand.strip()
                        if b_clean.lower() in target_set_lower:
                            # This is a target brand. Use its original input name if possible
                            # To keep the user's spelling/casing
                            final_analysis[brand] = data
                        else:
                            # Aggregate into Others
                            mentions = data.get("mentions", 0)
                            if mentions > 0:
                                others_agg["mentions"] += mentions
                                others_agg["articles"] += data.get("articles", 0)
                                
                                for src, count in data.get("sources", {}).items():
                                    others_agg["sources"][src] = others_agg["sources"].get(src, 0) + count
                                    
                                for dt, count in data.get("timeline", {}).items():
                                    others_agg["timeline"][dt] = others_agg["timeline"].get(dt, 0) + count
                                    
                                for s_type, s_count in data.get("sentiment", {}).items():
                                    others_agg["sentiment"][s_type] += s_count
                                    
                                # Merge some samples
                                for s_type, samples in data.get("article_samples", {}).items():
                                    if len(others_agg["article_samples"][s_type]) < 10:
                                        others_agg["article_samples"][s_type].extend(samples[:2])

                    if others_agg["mentions"] >= 0: # Change to >= 0 to always show
                        final_analysis["Others"] = others_agg

                    st.session_state.ca_analysis_results = final_analysis
                        
        if st.session_state.get('ca_analysis_results') is not None:
            analysis_results = st.session_state.ca_analysis_results
            brands_to_analyze = st.session_state.ca_brands_to_analyze
            
            st.markdown("---")

            st.markdown("### 📢 Share of Voice Overview")
            
            # Data prep for charts with Segregation
            pie_data = []
            timeline_data = []
            source_data = []
            sentiment_data = []
            
            # Define threshold for segregation
            all_mentions = [d.get("mentions", 0) for d in analysis_results.values() if d.get("mentions", 0) > 0]
            # Use median or average? Median is better for outliers. 
            threshold = pd.Series(all_mentions).median() if all_mentions else 0
            
            for brand, data in analysis_results.items():
                mentions = data.get("mentions", 0)
                
                # Segregation logic
                if brand == "Others":
                    cat = "Major Global Firms"
                else:
                    cat = "Major Global Firms" if mentions >= threshold else "Emerging & Niche Entities"
                
                # Determine if this is a brand explicitly requested by the user
                is_target = brand in brands_to_analyze or brand == "Others"
                
                if mentions >= 0:
                    if brand == "Others" and mentions == 0:
                        pass # Keep in summary but don't plot if 0
                    else:
                        pie_data.append({"Brand": brand, "Mentions": mentions, "Category": cat})
                    
                    # 1. Timeline
                    tl = data.get("timeline", {})
                    if tl:
                        for dt, count in tl.items():
                            if dt != "Unknown":
                                timeline_data.append({"Brand": brand, "Date": dt, "Mentions": count, "Category": cat})
                    elif is_target:
                        # Placeholder if target has no timeline data
                        last_30_days = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
                        timeline_data.append({"Brand": brand, "Date": last_30_days, "Mentions": 0, "Category": cat})
                    
                    # 2. Sources
                    srcs = data.get("sources", {})
                    if srcs:
                        for src, count in srcs.items():
                            source_data.append({"Brand": brand, "Source": src, "Mentions": count, "Category": cat})
                    elif is_target:
                        # Placeholder if target has no source data
                        source_data.append({"Brand": brand, "Source": "None Detected", "Mentions": 0, "Category": cat})
                        
                    # 3. Sentiment PERCENTAGE
                    sent = data.get("sentiment", {})
                    total_s = sum(sent.values())
                    if total_s > 0:
                        for s_type, s_count in sent.items():
                            pct = (s_count / total_s) * 100
                            sentiment_data.append({"Brand": brand, "Sentiment": s_type, "Percentage": round(pct, 1), "Category": cat})
                    elif is_target:
                        # Placeholder if target has no sentiment data
                        for s_type in ["Positive", "Neutral", "Negative"]:
                            sentiment_data.append({"Brand": brand, "Sentiment": s_type, "Percentage": 0, "Category": cat})
            
            import plotly.express as px
            
            if not pie_data:
                st.warning("No mentions found for any of the requested brands. (0 exposure)")
            else:
                # 0. High Level Metrics Overview
                st.markdown("### 🏆 Mention Counts Summary")
                display_list = list(brands_to_analyze)
                # Always append Others to summary if checking the pool
                if "Others" not in display_list:
                    display_list.append("Others")
                
                # Show in columns (4 per row)
                cols = st.columns(4)
                for i, brand in enumerate(display_list):
                    data = analysis_results.get(brand, {})
                    with cols[i % 4]:
                        st.metric(label=brand, value=data.get("mentions", 0))
                
                st.markdown("---")
                
                # 0.5 Top 20 Indian Publications
                st.markdown("### 🇮🇳 Top 20 Publications (India Region)")
                indian_source_counts = {}
                for brand, data in analysis_results.items():
                    for src, count in data.get("sources", {}).items():
                        if is_indian_source(src):
                            indian_source_counts[src] = indian_source_counts.get(src, 0) + count
                
                if indian_source_counts:
                    top_10_in_df = pd.DataFrame([
                        {"Publication": s, "Mentions": c} for s, c in indian_source_counts.items()
                    ]).sort_values(by="Mentions", ascending=False).head(20)
                    
                    # Only show Publication names without index numbers
                    st.dataframe(top_10_in_df[["Publication"]], hide_index=True, use_container_width=True)
                else:
                    st.info("No Indian publication data detected in this analysis.")

                st.markdown("---")
                
                # 1. Share of Voice Segregated
                st.markdown("### 📊 Market Share Analysis")
                for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                    cat_data = [d for d in pie_data if d["Category"] == cat]
                    if cat_data:
                        st.markdown(f"#### {cat}")
                        ov_type = st.selectbox("Select Visualization Type", ["Bar Chart", "Pie Chart"], key=f"ov_type_{cat}")
                        
                        if ov_type == "Bar Chart":
                            fig_ov = px.bar(cat_data, x="Brand", y="Mentions", color="Brand", 
                                            title=f"Exposure Share: {cat}", text="Mentions",
                                            template=get_chart_theme())
                            fig_ov.update_traces(texttemplate='%{text}', textposition='outside')
                        else:
                            fig_ov = px.pie(cat_data, names="Brand", values="Mentions", 
                                            title=f"Exposure Share: {cat}", hole=0.4,
                                            template=get_chart_theme())
                            fig_ov.update_traces(textinfo='percent+label')

                        fig_ov.update_layout(
                            paper_bgcolor="rgba(0,0,0,0)", 
                            plot_bgcolor="rgba(0,0,0,0)", 
                            font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"),
                            xaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF")),
                            yaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"))
                        )
                        st.plotly_chart(fig_ov, use_container_width=True)
                
                st.markdown("---")
                
                # 2. Timeline Segregated
                st.markdown("### 📈 Exposure Trends Over Time")
                for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                    cat_timeline = [d for d in timeline_data if d["Category"] == cat]
                    if cat_timeline:
                        st.markdown(f"#### {cat}")
                        cat_timeline.sort(key=lambda x: x["Date"])
                        tl_type = st.selectbox("Timeline Chart Type", ["Line Chart", "Bar Chart", "Area Chart"], key=f"tl_type_{cat}")
                        
                        if tl_type == "Line Chart":
                            fig_line = px.line(cat_timeline, x="Date", y="Mentions", color="Brand", markers=True, title=f"Trend: {cat}", template=get_chart_theme())
                        elif tl_type == "Bar Chart":
                            fig_line = px.bar(cat_timeline, x="Date", y="Mentions", color="Brand", barmode="group", title=f"Trend: {cat}", template=get_chart_theme())
                        else:
                            fig_line = px.area(cat_timeline, x="Date", y="Mentions", color="Brand", title=f"Trend: {cat}", template=get_chart_theme())

                        fig_line.update_layout(
                            paper_bgcolor="rgba(0,0,0,0)", 
                            plot_bgcolor="rgba(0,0,0,0)", 
                            font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"),
                            xaxis=dict(
                                title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"), 
                                tickfont=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"),
                                dtick="604800000", # 7 days in ms
                                tickformat="%b %d"
                            ),
                            yaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"), tickfont=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"))
                        )
                        st.plotly_chart(fig_line, use_container_width=True)
                
                st.markdown("---")
                
                # 2.5 Media Sources Analysis Segregated
                st.markdown("### 📰 Media Source Distribution")
                for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                    cat_src = [d for d in source_data if d["Category"] == cat]
                    if cat_src:
                        st.markdown(f"#### {cat}")
                        # Filter to top 50 sources to avoid clutter
                        cat_src.sort(key=lambda x: x["Mentions"], reverse=True)
                        fig_src = px.bar(cat_src[:50], x="Brand", y="Mentions", color="Source", 
                                         title=f"Source Distribution (Top 50): {cat}", barmode="stack",
                                         template=get_chart_theme())
                        fig_src.update_layout(
                            paper_bgcolor="rgba(0,0,0,0)", 
                            plot_bgcolor="rgba(0,0,0,0)", 
                            font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"),
                            xaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF")),
                            yaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"))
                        )
                        st.plotly_chart(fig_src, use_container_width=True)
                
                st.markdown("---")
                # 3. Sentiment PERCENTAGE Segregated
                st.markdown("### 🎭 Sentiment Landscape (%)")
                if sentiment_data:
                    for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                        cat_sent = [d for d in sentiment_data if d["Category"] == cat]
                        if cat_sent:
                            st.markdown(f"#### {cat}")
                            fig_sent = px.bar(cat_sent, x="Brand", y="Percentage", color="Sentiment", 
                                              title=f"Sentiment Distribution (%): {cat}",
                                              color_discrete_map={"Positive": "#27ae60", "Neutral": "#bdc3c7", "Negative": "#e74c3c"},
                                              barmode="group", text="Percentage",
                                              template=get_chart_theme())
                            fig_sent.update_traces(texttemplate='%{text}%', textposition='outside')
                            fig_sent.update_layout(
                                paper_bgcolor="rgba(0,0,0,0)", 
                                plot_bgcolor="rgba(0,0,0,0)", 
                                font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"),
                                yaxis_title="Percentage (%)",
                                xaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF")),
                                yaxis=dict(title_font=dict(color="#000000" if st.session_state.theme == 'light' else "#FFFFFF"))
                            )
                            st.plotly_chart(fig_sent, use_container_width=True)
                    
                    # DRILL DOWN - Inside the 'if sentiment_data' block but outside the 'cat' loop
                    st.markdown("---")
                    with st.expander("🔍 Explore Articles by Sentiment"):
                        st.markdown("Select a brand and sentiment to see the specific articles driving these metrics.")
                        
                        col_sed1, col_sed2 = st.columns(2)
                        with col_sed1:
                            drill_brands = list(brands_to_analyze)
                            if "Others" in analysis_results:
                                drill_brands.append("Others")
                            drill_brand = st.selectbox("Select Brand", drill_brands, key="drill_brand")
                        with col_sed2:
                            drill_sent = st.selectbox("Select Sentiment", ["Positive", "Neutral", "Negative"], key="drill_sent")
                        
                        # Retrieve the saved sample articles
                        brand_data = analysis_results.get(drill_brand, {})
                        article_samples = brand_data.get("article_samples", {})
                        selected_articles = article_samples.get(drill_sent, [])
                        
                        if selected_articles:
                            df_samples = pd.DataFrame(selected_articles)
                            st.dataframe(
                                df_samples,
                                column_config={
                                    "url": st.column_config.LinkColumn("Article Link"),
                                    "title": "Headline",
                                    "source": "Publication",
                                    "published": "Date"
                                },
                                hide_index=True,
                                use_container_width=True
                            )
                        else:
                            st.info(f"No {drill_sent} articles found for {drill_brand}.")
                else:
                    st.info("Sentiment data not available or all mentions are Neutral.")
                    
                st.markdown("---")
                
                st.markdown("---")
                # --- REPORT CONFIGURATION SECTION ---
                st.subheader("🛠️ Finalize Report Content")
                
                # Checkbox View (Toggle via state or always show if analysis exists)
                with st.expander("📝 Select Charts & Tables to Include in Export", expanded=True):
                    col_cfg1, col_cfg2 = st.columns(2)
                    with col_cfg1:
                        inc_metrics = st.checkbox("Brand Metrics Table", value=True)
                        inc_ov = st.checkbox("Overall Mention Share (Pie)", value=True)
                        inc_line = st.checkbox("Mentions Over Time (Line)", value=True)
                    with col_cfg2:
                        inc_top_pub = st.checkbox("Top 20 Indian Publications", value=True)
                        inc_bar = st.checkbox("Mentions by Publication (Per Brand)", value=True)
                        inc_sent = st.checkbox("Sentiment Analysis (Bar)", value=True)

                # --- DATA PREP FOR EXPORT ---
                metrics_list = []
                for brand, data in analysis_results.items():
                    metrics_list.append({
                        "Brand": brand,
                        "Mentions": data.get("mentions", 0),
                        "Articles": data.get("articles", 0),
                        "Avg Mentions/Article": round(data.get("mentions", 0) / data.get("articles", 1), 2) if data.get("articles", 0) > 0 else 0
                    })
                metrics_df = pd.DataFrame(metrics_list).sort_values(by="Mentions", ascending=False)
                
                if inc_metrics:
                    st.markdown("### 📊 Metrics Summary")
                    st.table(metrics_df)

                st.markdown("<br>", unsafe_allow_html=True)
                export_col1, export_col2 = st.columns(2)

                # --- EXPORT TO WORD (.docx) ---
                with export_col1:
                    try:
                        from docx import Document
                        from docx.shared import Inches
                        
                        doc = Document()
                        doc.add_heading('Brand Analysis Report', 0)
                        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                        doc.add_paragraph(f"Analyzed Brands: {', '.join(brands_to_analyze)}")
                        
                        def add_fig_to_doc(fig_obj, title):
                            if fig_obj:
                                # Force professional bright template for Word documents
                                fig_obj.update_layout(template="plotly_white", paper_bgcolor="white", plot_bgcolor="white", font=dict(color="black"))
                                img_bytes = fig_obj.to_image(format="png", width=800, height=500, scale=3)
                                doc.add_heading(title, level=2)
                                doc.add_picture(BytesIO(img_bytes), width=Inches(6.0))
                                doc.add_page_break()

                        def add_fig_html(fig_obj, title):
                            if fig_obj:
                                # Force professional bright template for HTML reports
                                fig_obj.update_layout(template="plotly_white", paper_bgcolor="white", plot_bgcolor="white", font=dict(color="black"))
                                div_html = fig_obj.to_html(full_html=False, include_plotlyjs=False)
                                return f"<div class='section-title'><h2>{title}</h2></div><div class='chart-container'>{div_html}</div>"
                            return ""

                        if inc_metrics:
                            doc.add_heading("Raw Exposure Metrics", level=2)
                            table = doc.add_table(rows=1, cols=len(metrics_df.columns))
                            table.style = 'Table Grid'
                            for i, col in enumerate(metrics_df.columns): table.rows[0].cells[i].text = col
                            for _, row in metrics_df.iterrows():
                                row_cells = table.add_row().cells
                                for i, val in enumerate(row): row_cells[i].text = str(val)
                            doc.add_page_break()

                        if inc_top_pub:
                            indian_source_counts = {}
                            for brand, data in analysis_results.items():
                                for src, count in data.get("sources", {}).items():
                                    if is_indian_source(src):
                                        indian_source_counts[src] = indian_source_counts.get(src, 0) + count
                            if indian_source_counts:
                                top_10_in_df = pd.DataFrame([{"Publication": s, "Mentions": c} for s, c in indian_source_counts.items()]).sort_values(by="Mentions", ascending=False).head(20)
                                doc.add_heading("Top 20 Indian Publications", level=2)
                                table_in = doc.add_table(rows=1, cols=1)
                                table_in.style = 'Table Grid'
                                table_in.rows[0].cells[0].text = "Publication"
                                for _, row in top_10_in_df.iterrows():
                                    row_cells = table_in.add_row().cells
                                    row_cells[0].text = str(row['Publication'])
                                doc.add_page_break()

                        # 11. Add Segregated Charts to Word
                        for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                            if inc_ov:
                                cat_pie = [d for d in pie_data if d["Category"] == cat]
                                if cat_pie:
                                    fig_cat_ov = px.bar(cat_pie, x="Brand", y="Mentions", color="Brand", title=f"Exposure Share: {cat}")
                                    add_fig_to_doc(fig_cat_ov, f"Market Share: {cat}")
                            
                            if inc_line:
                                cat_timeline = [d for d in timeline_data if d["Category"] == cat]
                                if cat_timeline:
                                    cat_timeline.sort(key=lambda x: x["Date"])
                                    fig_cat_line = px.line(cat_timeline, x="Date", y="Mentions", color="Brand", markers=True, title=f"Trend: {cat}")
                                    add_fig_to_doc(fig_cat_line, f"Exposure Timeline: {cat}")
                            
                            if inc_bar:
                                cat_src = [d for d in source_data if d["Category"] == cat]
                                if cat_src:
                                    cat_src.sort(key=lambda x: x["Mentions"], reverse=True)
                                    fig_cat_src = px.bar(cat_src[:50], x="Brand", y="Mentions", color="Source", 
                                                     title=f"Source Distribution: {cat}", barmode="stack")
                                    add_fig_to_doc(fig_cat_src, f"Media Sources: {cat}")

                            if inc_sent:
                                cat_sent = [d for d in sentiment_data if d["Category"] == cat]
                                if cat_sent:
                                    fig_cat_sent = px.bar(cat_sent, x="Brand", y="Percentage", color="Sentiment", 
                                                      title=f"Sentiment Distribution (%): {cat}",
                                                      color_discrete_map={"Positive": "#27ae60", "Neutral": "#bdc3c7", "Negative": "#e74c3c"},
                                                      barmode="group")
                                    add_fig_to_doc(fig_cat_sent, f"Sentiment Analysis: {cat}")
                        
                        doc_buffer = BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        st.download_button(
                            label="📄 Download Word Report (.docx)",
                            data=doc_buffer,
                            file_name="brand_analysis_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Word Error: {e}")

                # --- EXPORT TO WEB DOC (.html) ---
                with export_col2:
                    try:
                        metrics_table_html = ""
                        if inc_metrics:
                            metrics_table_html = "<h2 class='section-title'>Exposure Metrics</h2><table style='width:100%; border-collapse:collapse; margin:20px 0;'><thead><tr style='background:#2c3e50; color:white; text-align:left;'><th style='padding:10px;'>Brand</th><th style='padding:10px;'>Mentions</th><th style='padding:10px;'>Articles</th><th style='padding:10px;'>Avg</th></tr></thead><tbody>"
                            for _, row in metrics_df.iterrows():
                                metrics_table_html += f"<tr style='border-bottom:1px solid #eee;'><td style='padding:10px;'>{row['Brand']}</td><td style='padding:10px;'>{row['Mentions']}</td><td style='padding:10px;'>{row['Articles']}</td><td style='padding:10px;'>{row['Avg Mentions/Article']}</td></tr>"
                            metrics_table_html += "</tbody></table>"

                        html_content = f"""
                        <html><head><title>Brand Report</title><script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
                        <style>
                            body {{ font-family: Arial; padding: 40px; background: #f0f2f5; }}
                            .page {{ background: white; max-width: 900px; margin: auto; padding: 50px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }}
                            .section-title {{ border-left: 5px solid #3498db; padding-left: 15px; margin: 30px 0; color: #2c3e50; }}
                        </style></head>
                        <body><div class="page"><h1>Brand Intelligence Report</h1>{metrics_table_html}"""
                        
                        if inc_top_pub:
                            indian_source_counts = {}
                            for brand, data in analysis_results.items():
                                for src, count in data.get("sources", {}).items():
                                    if is_indian_source(src):
                                        indian_source_counts[src] = indian_source_counts.get(src, 0) + count
                            if indian_source_counts:
                                top_10_in_df = pd.DataFrame([{"Publication": s, "Mentions": c} for s, c in indian_source_counts.items()]).sort_values(by="Mentions", ascending=False).head(20)
                                html_content += "<div class='section-title'><h2>Top 20 Indian Publications</h2></div><table style='width:100%; border-collapse:collapse; margin:20px 0;'><thead><tr style='background:#f39c12; color:white; text-align:left;'><th style='padding:10px;'>Publication</th></tr></thead><tbody>"
                                for _, row in top_10_in_df.iterrows():
                                    html_content += f"<tr style='border-bottom:1px solid #eee;'><td style='padding:10px;'>{row['Publication']}</td></tr>"
                                html_content += "</tbody></table>"

                        # 11. Add Segregated Charts to HTML
                        for cat in ["Major Global Firms", "Emerging & Niche Entities"]:
                            # Share of Voice
                            if inc_ov:
                                cat_pie = [d for d in pie_data if d["Category"] == cat]
                                if cat_pie:
                                    fig_cat_ov = px.bar(cat_pie, x="Brand", y="Mentions", color="Brand", title=f"Exposure Share: {cat}")
                                    html_content += add_fig_html(fig_cat_ov, f"Market Share: {cat}")
                            
                            # Timeline
                            if inc_line:
                                cat_timeline = [d for d in timeline_data if d["Category"] == cat]
                                if cat_timeline:
                                    cat_timeline.sort(key=lambda x: x["Date"])
                                    fig_cat_line = px.line(cat_timeline, x="Date", y="Mentions", color="Brand", markers=True, title=f"Trend: {cat}")
                                    html_content += add_fig_html(fig_cat_line, f"Exposure Timeline: {cat}")
                            
                            if inc_bar:
                                cat_src = [d for d in source_data if d["Category"] == cat]
                                if cat_src:
                                    cat_src.sort(key=lambda x: x["Mentions"], reverse=True)
                                    fig_cat_src = px.bar(cat_src[:50], x="Brand", y="Mentions", color="Source", 
                                                     title=f"Source Distribution: {cat}", barmode="stack")
                                    html_content += add_fig_html(fig_cat_src, f"Media Sources: {cat}")

                            # Sentiment
                            if inc_sent:
                                cat_sent = [d for d in sentiment_data if d["Category"] == cat]
                                if cat_sent:
                                    fig_cat_sent = px.bar(cat_sent, x="Brand", y="Percentage", color="Sentiment", 
                                                      title=f"Sentiment landscape (%): {cat}",
                                                      color_discrete_map={"Positive": "#27ae60", "Neutral": "#bdc3c7", "Negative": "#e74c3c"},
                                                      barmode="group")
                                    html_content += add_fig_html(fig_cat_sent, f"Sentiment Analysis: {cat}")
                        
                        html_content += "</div></body></html>"
                        
                        st.download_button(
                            label="🌐 Download Web Report (.html)",
                            data=html_content,
                            file_name="brand_report.html",
                            mime="text/html",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"HTML Error: {e}")    
    # --- DISPLAY RESULTS (PAGINATED DB) ---

    st.subheader(f"📋 Results")
    
    # Filter Controls
    col_filter1, col_filter2 = st.columns([2, 2])
    with col_filter1:
        search_filter = st.text_input("🔍 Filter by Keyword", placeholder="Search in database...")
    with col_filter2:
        sector_filter = st.selectbox("📂 Filter by Sector", ["All", "Technology", "Artificial Intelligence", "Finance", "Health", "Sustainability", "Startups"])
        
    # Stats
    total_count = db_manager.get_total_count(sector_filter, search_filter)
    
    if total_count == 0:
        st.info("No articles found in the database matching your filters.")
    else:
        # Pagination Controls
        PAGE_SIZE = 50
        total_pages = (total_count // PAGE_SIZE) + (1 if total_count % PAGE_SIZE > 0 else 0)
        
        col_page1, col_page2, col_page3 = st.columns([1, 2, 1])
        with col_page2:
            current_page = st.number_input("Page", min_value=1, max_value=max(1, total_pages), value=1)
            
        st.write(f"Showing page **{current_page}** of **{total_pages}** ({total_count} total articles)")
        
        offset = (current_page - 1) * PAGE_SIZE
        
        # Fetch ONLY the current page
        articles_to_display = db_manager.get_articles(
            limit=PAGE_SIZE, 
            offset=offset, 
            sector_filter=sector_filter, 
            search_query=search_filter
        )
        
        # --- SCROLLABLE CONTAINER ---
        with st.container(height=800):
            for i, article in enumerate(articles_to_display):
                title = article['title']
                source = article['source']
                summary = article.get('summary', 'No summary.')
                full_text = article.get('full_text', '')
                link = article['link']
                published = article['published']
                
                st.markdown(f"### {offset + i + 1}. [{title}]({link})")
                st.caption(f"**Source:** {source} | **Published:** {published}")
                
                with st.expander("📖 Read Full Article Content"):
                    st.markdown("#### Summary")
                    st.info(summary)
                    
                    st.markdown("#### Full Article")
                    st.write(full_text if full_text else "⚠️ Could not extract full text.")

                st.markdown(f"🔗 [**Go to Article**]({link})")
                st.markdown("---")

        # --- DOWNLOAD BUTTONS ---
        # Fetch ALL matching articles for download (not just current page)
        # Warning: This might be heavy if >10k articles, but better than partial download
        if total_count > 0:
            st.write("Preparing download...")
            all_articles_df = pd.DataFrame(db_manager.get_articles(limit=10000, offset=0, sector_filter=sector_filter, search_query=search_filter))
            
            # Prepare Excel file
            buffer = BytesIO()
            with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                all_articles_df.to_excel(writer, index=False, sheet_name='News')
            buffer.seek(0)
            
            col_dl1, col_dl2, col_dl3 = st.columns(3)
            with col_dl1:
                st.download_button(
                    label="📥 Download as Excel",
                    data=buffer,
                    file_name=f"news_results.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with col_dl2:
                csv = all_articles_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download as CSV",
                    data=csv,
                    file_name=f"news_results.csv",
                    mime="text/csv"
                )
            with col_dl3:
                # Add download for the entire database as a SQL dump
                sql_dump = db_manager.get_db_sql_dump()
                st.download_button(
                    label="📥 Download SQL Dump",
                    data=sql_dump,
                    file_name="articles_dump.sql",
                    mime="application/sql"
                )

# --- APP FLOW ---

# Initialize session state for auth
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'username' not in st.session_state:
    st.session_state['username'] = None

# Logic to switch between Login/Signup and Main App
if not st.session_state['logged_in']:
    tab1, tab2 = st.tabs(["Login", "Sign Up"])
    
    with tab1:
        login_page()
    
    with tab2:
        signup_page()
else:
    show_dashboard()


